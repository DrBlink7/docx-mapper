import re
import docx
import openpyxl
from datetime import datetime

# Flag to control date formatting
USE_DATETIME = True  # Set to False to format dates as 'dd/mm/yyyy hh:mm:ss'

def load_mapping(file_path):
    """
    Load the Excel file and build the mapping dictionary.
    """
    try:
        workbook = openpyxl.load_workbook(file_path)
        sheet = workbook.active
        mapping = {}
        for row in sheet.iter_rows(min_row=1, values_only=True):
            key, value = row
            if key:
                # Format the date if the flag is active
                if USE_DATETIME and isinstance(value, datetime):
                    mapping[key] = value.strftime('%d/%m/%Y')
                else:
                    mapping[key] = value
        return mapping
    except FileNotFoundError:
        raise FileNotFoundError(f"Excel file '{file_path}' not found.")
    except Exception as e:
        raise Exception(f"Error loading mapping: {e}")

def replace_text_in_paragraph(paragraph, mapping):
    """
    Replace placeholders in a paragraph with corresponding values from the mapping.
    """
    runs = paragraph.runs
    full_text = ''.join(run.text for run in runs)
    updated_text = full_text
    for key, value in mapping.items():
        updated_text = re.sub(r'\{\{' + re.escape(key) + r'\}\}', str(value), updated_text)
    if runs:
        runs[0].text = updated_text
        for run in runs[1:]:
            run.text = ''

def process_document(doc_path, mapping, output_path):
    """
    Process the Word document to replace placeholders with mapping values.
    """
    try:
        document = docx.Document(doc_path)
        # Process paragraphs
        for paragraph in document.paragraphs:
            replace_text_in_paragraph(paragraph, mapping)
        # Process tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, mapping)
        # Save the modified document
        document.save(output_path)
    except FileNotFoundError:
        raise FileNotFoundError(f"Word document '{doc_path}' not found.")
    except Exception as e:
        raise Exception(f"Error processing document: {e}")

def main():
    """
    Main function to execute the script.
    """
    excel_file = 'mapping.xlsx'
    word_file = 'base_document.docx'
    output_file = 'final_document.docx'

    try:
        # Load mapping from Excel
        mapping = load_mapping(excel_file)
        # Process the Word document
        process_document(word_file, mapping, output_file)
        print(f"Document processed successfully. Saved as '{output_file}'.")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    main()
